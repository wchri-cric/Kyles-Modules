import sys
sys.path.append("/usr/lib/python3.5/site-packages")

import os,json
from copy import deepcopy
from math import floor, ceil
from datetime import datetime as dt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import table
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK,WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from random import randrange as rand

FILEPATH = os.path.dirname(os.path.abspath(__file__))+"/"
MAKE_ID = lambda: "".join([str(hex(rand(16)))[2:].upper() for _ in range(8)])

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
REL = "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
OV = "{http://schemas.openxmlformats.org/package/2006/content-types}Override"

def checkedElement():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'),"true")
    return elm

# line_text: the text in the footer formatted line this: Initials _________
# sub_str: the string to add to the line: e.g initials or date
def addtoLine(line_text,sub_str):
    line_indexes = [i for i,x in enumerate(line_text) if x=="_"]
    line_start, line_end = line_indexes[0], line_indexes[-1]
    line = line_text[line_start:line_end]
    line_length = len(line) - len(sub_str)
    new_line = line[:floor(line_length/2)] + sub_str + line[ceil(line_length/2)+len(sub_str)-1:]
    return line_text[:line_start] + new_line + line_text[line_end:]

def fixDates(old_dates):
    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    new_dates = []
    for d in old_dates:
        d = dt.strptime(d,"%Y-%m-%d") if d else ""
        new_dates.append("-".join([str(d.day),months[d.month-1],str(d.year)]) if d else "")
    return new_dates

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def makeSections(table,N,length=0):
    if not len(N): return [table]
    n = N[0] - length
    new_table = deepcopy(table)
    # Remove trailing rows from table
    rows = table.rows[n+1:]
    for row in rows: remove_row(table, row)
    # Remove leading rows from new table
    rows = new_table.rows[:n+1]
    for row in rows: remove_row(new_table,row)

    new_tables = makeSections(new_table,N[1:],length+len(table.rows))
    return [table] + new_tables


if __name__ == "__main__":
    config_str = ""
    for i in range(1,len(sys.argv)): config_str += sys.argv[i]

    configs = json.loads(config_str)

    failed = configs["failed"]
    breaksteps = configs["breaksteps"]
    reference_text = configs["reference_text"]
    initials = configs["initials"]
    dates = fixDates(configs["dates"])
    redcap_version = configs["redcap_version"]
    browser_version = configs["browser_version"]
    os_version = configs["os_version"]
    notes = configs["notes"]
    name = configs["name"]
    environment = configs["environment"]

    # Get the number of sections in the input document
    pre_doc = Document(FILEPATH+"in.docx")
    additional_sections = len(pre_doc.sections)

    # Fix the dates list so that the first date will be applied to all the sections preceding the first break added to the document
    dates = [dates[0]] * additional_sections + dates[1:]
    doc = Document(FILEPATH+"in.docx")

    steps_table = doc.tables[0] # This is the table with the steps in it (the one that will be split)
    final_table = doc.tables[1] # This is the summary table on the last page

    ftbl = final_table._tbl
    ftbl = deepcopy(ftbl)
    final_table._element.getparent().remove(final_table._element) # Remove table

    # Remove the empty paragraphs (and possible page break) after the last table
    for i in range(len(doc.paragraphs)):
        if type(doc.paragraphs[i]._element.getprevious()) == table.CT_Tbl: break
    for _ in range(len(doc.paragraphs)-i): remove_paragraph(doc.paragraphs[-1])

    # Make the new tables to add (first table will already be in the document)
    sections = makeSections(steps_table,breaksteps)
    for i in range(1,len(sections)):
        s = doc.add_section(WD_SECTION.CONTINUOUS)
        p = doc.add_paragraph('')
        p._element.addnext(sections[i]._tbl)

    # Add a page break before final page
    p = doc.add_paragraph('')
    #p._element.addprevious(steps_table_copy._tbl)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p._element.addnext(ftbl) # Add final table

    # Unlink sections
    footer_table = doc.sections[0].footer.tables[0]
    for i in range(additional_sections,len(doc.sections)):
        footer = doc.sections[i].footer
        footer.is_linked_to_previous = False
        t = footer.add_table(1,2,Cm(40.63))
        for j in range(2): t.row_cells(0)[1].add_paragraph(footer_table.row_cells(0)[1].paragraphs[j].text)
        t.rows[0].witdh = Cm(1.13)
        t.row_cells(0)[1].width = Cm(3.65)
        remove_paragraph(footer.paragraphs[0])
        footer.add_paragraph('')

    # REDCap version
    doc.paragraphs[0].text = addtoLine(doc.paragraphs[0].text,redcap_version)
    doc.paragraphs[0].runs[-1].font.bold = False
    doc.paragraphs[0].runs[-1].font.size = Pt(10)

    # Environment
    env_text = doc.paragraphs[0].runs[-1].text
    previous_text = env_text[:env_text.index(environment)]
    next_text = env_text[env_text.index(environment)+len(environment):]
    doc.paragraphs[0].runs[-1].text = previous_text
    doc.paragraphs[0].runs[-1].font.bold = False
    doc.paragraphs[0].runs[-1].font.size = Pt(10)
    doc.paragraphs[0].add_run(environment)
    doc.paragraphs[0].runs[-1].font.bold = True
    doc.paragraphs[0].runs[-1].font.underline = True
    doc.paragraphs[0].runs[-1].font.size = Pt(10)
    doc.paragraphs[0].add_run(next_text)
    doc.paragraphs[0].runs[-1].font.bold = False
    doc.paragraphs[0].runs[-1].font.size = Pt(10)

    ind = doc.paragraphs[2].text.index("OS")
    browser_text = addtoLine(doc.paragraphs[2].text[:ind],browser_version)
    os_text = addtoLine(doc.paragraphs[2].text[ind:],os_version)

    # Browser Version
    doc.paragraphs[2].text = browser_text
    doc.paragraphs[2].runs[-1].font.bold = False
    doc.paragraphs[2].runs[-1].font.size = Pt(10)

    # OS Version
    doc.paragraphs[2].add_run(os_text)
    doc.paragraphs[2].runs[-1].font.bold = False
    doc.paragraphs[2].runs[-1].font.size = Pt(10)

    # Add footer initials and dates
    sections = doc.sections
    for i in range(additional_sections-1,len(sections)):
        footer = sections[i].footer
        initial_line = footer.tables[0].row_cells(0)[1].paragraphs[-2]
        date_line = footer.tables[0].row_cells(0)[1].paragraphs[-1]
        initial_line.text = addtoLine(initial_line.text, initials)
        date_line.text = addtoLine(date_line.text, dates[i])

    # Add checkboxes and reference text
    rowcount = 1
    for i in range(len(doc.tables)-1):
        steps = doc.tables[i]
        for j in range(1-bool(i),len(steps.rows)): # Skip first row in first table (since first row is column header)
            references = steps.row_cells(j)[-1]
            x = steps.row_cells(j)[-2]._element.xpath('.//w:checkBox')
            try:
                x[rowcount in failed].append(checkedElement())
            except IndexError:
                print("It seems that step #{} is missing one or more checkboxes".format(rowcount))
                quit()
            if str(rowcount) in reference_text.keys():
                if "___" in references.text: # Fill in line in references column (and keep text boldness)
                    is_bold = references.paragraphs[0].runs[-1].bold
                    references.text = addtoLine(references.text,reference_text[str(rowcount)])
                    references.paragraphs[0].runs[-1].bold = is_bold
                else:
                    references.add_paragraph(reference_text[str(rowcount)])
            rowcount +=1

    complete = doc.tables[-1]
    # Add notes
    p = complete.row_cells(1)[0].add_paragraph(notes)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add name and date
    complete.row_cells(3)[1].add_paragraph(name)
    complete.row_cells(3)[3].add_paragraph(dates[-1])

    # Bold "Yes" if no failed steps, bold "No" otherwise
    run = complete.row_cells(2)[2+bool(len(failed))].paragraphs[0].runs[0]
    run.font.underline = True
    run.font.bold = True

    doc.save(FILEPATH+"out.docx")
    print("Document Created Successfully")
